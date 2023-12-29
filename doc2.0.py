from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import locale
from docx.shared import Inches


def formatar_num(numero):
    # Definir a localização para o formato monetário desejado
    locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')

    # Formatando o número
    numero_formatado = locale.currency(numero, grouping=True, symbol=True)

    return numero_formatado

def atualiza_data():
    data_atual = datetime.now().date().strftime("%d-%m-%y")
    return data_atual
def hora_exata():
    hora_atual = datetime.now().strftime("%H:%M:%S")
    return hora_atual
nova_data = atualiza_data()
nova_hora = hora_exata()

def formata_data(nova_data):
    nova_data = nova_data.split('-')
    dic = {1: 'Janeiro',2:'Fevereiro',3: 'Março',4:"Abril", 5:"Maio", 6:"Junho", 7: 'Julho', 8:"Agosto", 9:"Setembro", 10:"Outubro", 11:"Novembro", 12:"Dezembro"}
    dia = (nova_data[0])
    mes = int(nova_data[1])
    ano = int(nova_data[2])
    if mes in dic:
        for chave, valor in dic.items():
            if chave == mes:
                return (f'{dia} de {valor} de 20{ano}')

def formatar_variaveis_negrito(texto):
    # Criar um documento temporário para acessar as funcionalidades do docx
    documento_temporario = Document()

    # Adicionar um parágrafo ao documento temporário
    paragrafo_temporario = documento_temporario.add_paragraph(texto)

    # Selecionar as variáveis entre {} e torná-las em negrito
    for variavel in paragrafo_temporario.runs:
        texto_variavel = variavel.text
        if texto_variavel.startswith('{') and texto_variavel.endswith('}'):
            variavel.bold = True

    # Obter o texto formatado com as variáveis em negrito
    texto_formatado = ""
    for run in paragrafo_temporario.runs:
        texto_formatado += run.text

    return texto_formatado


#Looping calculo bf
fim_bf = False
while not fim_bf:
    #Entrada de dados/correcao de virgula
    nome_projeto = input('Qual o nome do projeto\nDIGITE (apenas um para ser o titulo): \nexemplo de preenchimento: rua javri\n').upper()
    nome_projeto = nome_projeto.replace("/", "-")
    area_terreno = float(input('\nInsira a AREA DO TERRENO em m²: ').replace(',','.'))
    area_computavel = float(input('Insira a AREA COMPUTÁVEL em m²: ').replace(',','.'))
    valor_referencia = float(input('Insira o VALOR REFERENCIA: ').replace(',','.'))
    zona = int(input(f"Insira qual a zona desejada\nZona de Qualificação (1)\nZona de Reestruturação (2)\n"))
    
    # Função para calcular a fórmula desejada
    #varaiveis referentes ao processamento:
    cp = area_computavel/area_terreno
    fator_reducao = 0.8
    fmp = 5.0578    
    if zona == 1:
        ic = 0.4
        coeficiente_basico = 2.5
        cpc = round(cp - coeficiente_basico,2)
        bf = (area_terreno * valor_referencia * cpc * ic * fator_reducao) * fmp
        dados_bf = f'''
            ÁREA DO TERRENO = {('{:,.2f}'.format(area_terreno).replace(',', '.'))} m²
            ÁREA COMPUTAVEL = {('{:,.2f}'.format(area_computavel).replace(',', '.'))} m²
            VALOR DE REFERENCIA = {(valor_referencia):.2f} FMP/m²
            COEFICIENTE BASICO = {(coeficiente_basico):.2f}
            COEFICIENTE PROJETO = {(cp):.2f}
            COEFICINTE PRETENDIDO = {(cp):.2f} - {(coeficiente_basico):.2f} = {(cpc):.2f} = Cp
            FMP = R${(fmp)}

            Contrapartida financeira:
            Bf = At x Vr x Cp x Ic x Fr
            Bf = {'{:,.2f}'.format(area_terreno).replace(',', '.')} x {valor_referencia:.2f} x {cp:.2f} x {ic:.2f} x {fator_reducao:.2f} x {fmp} = 
        '''
    if zona == 2:
        ic = 0.33
        coeficiente_projeto = 3.0
        cpc = round(cp - coeficiente_projeto,2)
        bf = (area_terreno * valor_referencia * cpc * ic * fator_reducao) * fmp
        dados_bf = f'''
            ÁREA DO TERRENO = {('{:,.2f}'.format(area_terreno).replace(',', '.'))} m²
            ÁREA COMPUTAVEL = {('{:,.2f}'.format(area_computavel).replace(',', '.'))} m²
            VALOR DE REFERENCIA = {(valor_referencia):.2f} FMP/m²
            COEFICIENTE BASICO = {(cp):.2f}
            COEFICIENTE PROJETO = {(coeficiente_projeto):.2f}
            COEFICINTE PRETENDIDO = {(cp):.2f} - {(coeficiente_projeto):.2f} = {(cpc):.2f} = Cp
            FMP = R${(fmp)}

            Contrapartida financeira:
            Bf = At x Vr x Cp x Ic x Fr
            Bf = {'{:,.2f}'.format(area_terreno).replace(',', '.')} x {valor_referencia:.2f} x {cp:.2f} x {ic:.2f} x {fator_reducao:.2f} x {fmp}=
        '''

        
    elif cp > coeficiente_basico or cp > coeficiente_projeto:
        resultado_bf = f'RESULTADO OBTIDO:  {formatar_num(bf)};  CPC = {(cp):.2f}'
    elif cp < coeficiente_basico or coeficiente_projeto:
        resultado_bf = f'Resultado do CP menor que 2.5, Não é necessario pagar a ODC!; CPC = {(cp):.2f}'
    else:
        print('\nERRO! Zona invalida!!\n')


    encerrar = input('\nPara REFAZER o calculo insira "R"\nPara CONFIRMAR o calculo digite "C"\n').lower()
    if encerrar == ('r'):
        fim_bf = False
    if encerrar == ('c'):
        fim_bf = True

#lLooping calculo EIV
fim_eiv = False
while not fim_eiv:
    #tipo 1
    area_a_construir = float(input('\nInsira a Área a Construir:\n').replace(',','.'))
    taxa_eiv = 0.025
    indice_CUB = (input('Deseja alterar o indice ou manter o de JUNHO DE 2023 (R$1.954,65)?\nInsira "M" para manter;\nPara alterar, insira o novo valor:\n').replace(',','.')).lower()
    nr_vr = valor_referencia
    # float(input('Qual o valor do multiplicador referente ao valor de referencia ( nº x FMP / m² = VR )\nQual o valor de nº?\n').replace(',','.'))
    vr_eiv = round(fmp * nr_vr, 2)
    if indice_CUB == 'm':
        indice_CUB = 1954.65
    if indice_CUB != 'mnbvlkj':
        indice_CUB = float(indice_CUB)
        
    #calculo:
    eiv = ((area_terreno * vr_eiv) + ((area_a_construir) * (indice_CUB))) * taxa_eiv
    dados_eiv = f'''
        ÁREA DO TERRENO(At) = {'{:,.2f}'.format(area_terreno).replace(',', '.')} m²
        ÁREA Á CONSTRUIR(Ac) = {'{:,.2f}'.format(area_a_construir).replace(',', '.')}m²
        VR = VALOR DE REFERENCIA = {nr_vr} FMP/m² = R$ {vr_eiv}
        TAXA EIV/RIT - TIPO 1 = 2,5%
        I_CUB_R8N - índice - CUB - R 8 - N - SINDUSCON = R$ {indice_CUB}
        FMP = R$ {fmp}

        Calculo:
        EIV = [(At x VR) + (Ac x I_CUB_R8N)] x 2,5%
        EIV = ({(round(area_terreno, 2))} * {(round(vr_eiv, 2))} + {(round(area_a_construir,2))} * {(round(indice_CUB,2))}) x 0,025
        EIV = [({(round(area_terreno * vr_eiv, 2))}) + ({(round(area_a_construir * indice_CUB, 2))})] x 0,025
    '''
    resultado_eiv = eiv
    encerrar_eiv = input('\nPara REFAZER o calculo insira "R"\nPara CONFIRMAR o calculo digite "C"\n').lower()
    if encerrar_eiv == ('r'):
        fim_eiv = False
    if encerrar_eiv == ('c'):
        fim_eiv = True


#Looping de criação do documento:
fim_pr = False
while not fim_pr:
   # Criar um doc
    documento = Document()

    # Configurar cabeçalho
    secao = documento.sections[0]
    header = secao.header

    # Adicionar o parágrafo no cabeçalho
    paragrafo = header.add_paragraph()

    # Adicionar o "Texto 1" na esquerda
    texto1 = paragrafo.add_run("MG Barone Jr\nEngº Civil\n\n\n\n\t\t\t\t\t\t\t\t\t      ")
    texto1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph_format = paragrafo.paragraph_format
    paragraph_format.line_spacing = Pt(10)
    # Adicionar a imagem na direita
    imagem_path = "barone_logo.jpg"
    largura_maxima = Inches(1.7)
    altura_maxima = Inches(0.89)

    run_imagem = paragrafo.add_run()
    run_imagem.add_picture(imagem_path, width=largura_maxima, height=altura_maxima)
    run_imagem.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


    header.add_paragraph("__________________________________________________________________________________________________________________")
    # Definir o espaçamento do cabeçalho (em polegadas)
    secao.header_distance = Inches(0.3)

    # Definir as margens para a seção principal
    secao = documento.sections[0]
    secao.left_margin = Inches(1)
    secao.right_margin = Inches(1)
    secao.top_margin = Inches(1)
    secao.bottom_margin = Inches(1)

    # titulo
    titulo_style = documento.styles.add_style("Titulo", WD_PARAGRAPH_ALIGNMENT.CENTER)
    fonte_titulo = titulo_style.font
    fonte_titulo.name = "Calibri"
    fonte_titulo.size = Pt(12)
    fonte_titulo.bold = True
    fonte_titulo.color.rgb = RGBColor(00, 00,00)  # Cor preta
    
    # Texto comum
    texto_style = documento.styles.add_style("Texto", WD_PARAGRAPH_ALIGNMENT.CENTER)
    fonte_texto = texto_style.font
    fonte_texto.name = "Calibri"
    fonte_texto.size = Pt(12)
    fonte_texto.bold = False
    fonte_texto.color.rgb = RGBColor(00,00,00)

    #texto Rodapé
    texto_rodape = documento.styles.add_style("Rodape", WD_PARAGRAPH_ALIGNMENT.CENTER)
    fonte_rodape = texto_rodape.font
    fonte_rodape.name = "Calibri"
    fonte_rodape.size = Pt(9)
    fonte_rodape.bold = False
    fonte_rodape.color.rgb = RGBColor(0,127,255)
    
    # Texto Negrito
    texto_style = documento.styles.add_style("Negrito", WD_PARAGRAPH_ALIGNMENT.CENTER)
    fonte_texto = texto_style.font
    fonte_texto.name = "Calibri"
    fonte_texto.size = Pt(12)
    fonte_texto.bold = True
    fonte_texto.color.rgb = RGBColor(00,00,00)

    
    # Adiciona um estilo de sombra ao documento
    sombra_style = documento.styles.add_style("Sombra", WD_PARAGRAPH_ALIGNMENT.CENTER)
    fonte_sombra = sombra_style.font
    fonte_sombra.name = "Calibri"
    fonte_sombra.size = Pt(12)
    fonte_sombra.bold = True
    fonte_sombra.color.rgb = RGBColor(0, 0, 0)


    #corrigindo formatação em negrito:
    dados_bf_negrito = formatar_variaveis_negrito(dados_bf)
    dados_eiv_negrito = formatar_variaveis_negrito(dados_eiv)



    # add titulo
    titulo = documento.add_paragraph(f"MEMORIAL DE CALCULOS BASICOS PARA OODC E EIV/RIT-TIPO I-LEI 9.924/16 PROJETO {nome_projeto}", style = "Titulo") 
    # Alinhe o parágrafo no centro
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # add paragrafos
    titulo_bf = documento.add_paragraph('1. CALCULO DO BENEFICIO FINANCEIRO - OUTORGA ONEROSA DO DIREITO DE CONSTRUIR:', style = "Negrito")
    texto_bf = documento.add_paragraph(f"{(dados_bf_negrito)}", style = "Texto")
    paragrafo = documento.add_paragraph(f"\t{resultado_bf}")
    paragrafo.add_run().bold = True
    paragrafo.style = sombra_style
    # Define a cor de fundo do parágrafo como laranja escuro
    paragrafo_para = paragrafo._element.get_or_add_pPr()
    shading_element = parse_xml(r'<w:shd {} w:fill="FFA500"/>'.format(nsdecls('w')))
    paragrafo_para.append(shading_element)
    #espaçamentos
    paragraph_format = titulo_bf.paragraph_format
    paragraph_format.line_spacing = Pt(12)
    paragraph_format = texto_bf.paragraph_format
    paragraph_format.line_spacing = Pt(12)

    # paragrafos
    titulo_eiv = documento.add_paragraph('2. CALCULO DO ESTUDO DE IMPACTO DE VIZINHANÇA E DE TRANSITO - TIPO 1:', style = "Negrito")
    texto_eiv = documento.add_paragraph(f"{(dados_eiv_negrito)}", style = "Texto")
    paragrafo = documento.add_paragraph(f"\tEIV =  {formatar_num(eiv)}")
    paragrafo.add_run().bold = True
    paragrafo.style = sombra_style
    # Define a cor de fundo do parágrafo como laranja escuro
    paragrafo_para = paragrafo._element.get_or_add_pPr()
    shading_element = parse_xml(r'<w:shd {} w:fill="FFA500"/>'.format(nsdecls('w')))
    paragrafo_para.append(shading_element)
    #espaçamentos
    paragraph_format = titulo_eiv.paragraph_format
    paragraph_format.line_spacing = Pt(12)
    paragraph_format = texto_eiv.paragraph_format
    paragraph_format.line_spacing = Pt(12)


    # espaço para assinatura
    data_lugar = documento.add_paragraph(f'Santo André, SP, {formata_data(nova_data)}\n')
    linha_para_assinatura = documento.add_paragraph('__________________________________')
    nome_assinatura = documento.add_paragraph('Miguel G. Barone Jr.\nEngº Civil')
    # alinhamento e espaçamentos
    paragraph_format = nome_assinatura.paragraph_format
    paragraph_format.line_spacing = Pt(12)
    data_lugar.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    linha_para_assinatura.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    nome_assinatura.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    # Rodapé
    footer = secao.footer
    paragrafo_rodape = footer.paragraphs[0]
    paragrafo_rodape.text = "________________________________________________________________________________________________________\n\t\t\tAVENIDA PORTUGAL, 397 – 4º ANDAR – SALA 406 – CENTRO – SANTO ANDRÉ – SP\n\t\tFONE: 11 - 4438.23.52 – Cel / WhatsApp : +55 11 - 9.9949.11.65 – e-mail: barone@terra.com.br"
    paragrafo_rodape.style = texto_rodape
    paragraph_format = paragrafo_rodape.paragraph_format
    paragraph_format.line_spacing = Pt(12)

        
    nome_arquivo = f"{nome_projeto}-BF_EIV-{nova_data}.docx"

    # Concatenar o caminho absoluto com o nome do arquivo
    caminho_completo = f"{nome_arquivo}"

    print(f"Arquivo '{nome_projeto}-BF_EIV-{nova_data}/{nova_hora}.docx' foi criado com sucesso!")
    try:
        documento.save(caminho_completo)
    except Exception as erro:
        # Imprimir o possível erro
        print("Ocorreu um erro:", erro)
        
    fim_pr = input("Digite:\nE -> Encerrar\nC -> Criar outro documento\n").lower()

    if fim_pr == 'e':
        fim_pr = True
    else:
        fim_pr = False